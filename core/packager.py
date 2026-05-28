"""
Agent Factory — 自动化打包器（Packager）V5.0
=============================================
将生成的智能体脚本打包为最终交付物。

交付物清单：
  app.py                      — 智能体主程序
  配置_API_Key.env             — 环境变量模板
  一键启动_Windows.bat         — Windows 双击启动
  一键启动_Mac.sh              — macOS/Linux 双击启动
  requirements.txt             — 依赖清单
  使用说明书.docx               — Word 说明书（默认交付文档）
  使用说明书.md                 — Markdown 说明书（纯文本版本）
  部署到飞书_机器人配置.yaml     — 飞书应用配置
  部署到Coze_导入文件.json      — Coze Bot 导入配置
  部署到Dify_工作流导入.yml     — Dify 工作流 DSL

设计原则：
  - 文件名直观无 emoji，专业可交付
  - 去品牌化：不出现 "Agent Factory 全自动生成" 等字样
  - 默认交付 Word 说明书，符合企业场景
  - 平台部署文件默认全量生成
"""

import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger("agent_factory.packager")

# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------
_BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = _BASE_DIR / "output_agents"
TEMPLATES_DIR = _BASE_DIR / "templates"

PLATFORM_TEMPLATES = {
    "feishu": "feishu_bot.yaml",
    "coze": "coze_bot.json",
    "dify": "dify_workflow.yml",
}


class PackageResult:
    """打包结果"""

    def __init__(self, success: bool, output_path: str = "", message: str = ""):
        self.success = success
        self.output_path = output_path
        self.message = message

    def to_dict(self) -> dict:
        return {
            "success": self.success,
            "output_path": self.output_path,
            "message": self.message,
        }

    def __repr__(self) -> str:
        status = "OK" if self.success else "FAIL"
        return f"<PackageResult {status} {self.output_path or self.message}>"


class AgentPackager:
    """
    V5.0 自动化打包器

    :param script_path:   智能体脚本路径
    :param delivery_type: 交付类型（exe / zip / web / docx）
    :param agent_name:    智能体名称
    :param agent_scenario: 场景描述（用于说明书填充）
    """

    def __init__(
        self,
        script_path: str,
        delivery_type: str = "zip",
        agent_name: str = "未命名智能体",
        agent_scenario: str = "通用任务",
        web_script_path: Optional[str] = None,
    ):
        self.script_path = Path(script_path).resolve()
        self.delivery_type = delivery_type
        self.agent_name = agent_name
        self.agent_scenario = agent_scenario
        self.web_script_path = Path(web_script_path).resolve() if web_script_path else None

        if not self.script_path.exists():
            raise FileNotFoundError(f"智能体脚本不存在：{self.script_path}")

        if self.delivery_type not in ("exe", "zip", "web", "docx"):
            logger.warning("不支持的 delivery_type '%s'，回退到 'zip'", delivery_type)
            self.delivery_type = "zip"

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # 主入口
    # ------------------------------------------------------------------
    def package(self) -> PackageResult:
        """执行打包流程"""
        dispatcher = {
            "exe": self._build_exe,
            "zip": self._build_zip,
            "web": self._build_web,
            "docx": self._build_docx_only,
        }
        handler = dispatcher.get(self.delivery_type, self._build_zip)
        return handler()

    # ------------------------------------------------------------------
    # exe — PyInstaller 打包
    # ------------------------------------------------------------------
    def _build_exe(self) -> PackageResult:
        """使用 PyInstaller 打包为单文件可执行文件"""
        try:
            import PyInstaller  # noqa: F401
        except ImportError:
            return PackageResult(
                success=False,
                message="缺少 PyInstaller。请运行：pip install pyinstaller",
            )

        logger.info("开始 PyInstaller 打包：%s", self.script_path)
        try:
            output_name = self.script_path.stem

            with tempfile.TemporaryDirectory(prefix="agent_factory_") as tmpdir:
                tmpdir_path = Path(tmpdir)
                dist_dir = tmpdir_path / "dist"

                result = subprocess.run(
                    [
                        sys.executable, "-m", "PyInstaller",
                        "--onefile",
                        "--name", output_name,
                        "--distpath", str(dist_dir),
                        "--workpath", str(tmpdir_path / "build"),
                        "--specpath", str(tmpdir_path),
                        "--noconfirm",
                        str(self.script_path),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=300,
                )

                if result.returncode != 0:
                    logger.error("PyInstaller 打包失败：%s", result.stderr[-500:])
                    return PackageResult(
                        success=False,
                        message=f"PyInstaller 打包失败：{result.stderr[-300:]}",
                    )

                exe_files = list(dist_dir.glob(output_name + "*"))
                if not exe_files:
                    return PackageResult(
                        success=False,
                        message="PyInstaller 打包完成但未找到输出文件",
                    )

                output_path = OUTPUT_DIR / exe_files[0].name
                shutil.copy2(exe_files[0], output_path)

                logger.info("EXE 打包完成 -> %s", output_path)
                return PackageResult(
                    success=True,
                    output_path=str(output_path),
                    message="独立可执行文件已生成",
                )

        except subprocess.TimeoutExpired:
            return PackageResult(success=False, message="PyInstaller 打包超时（>5分钟）")
        except Exception as e:
            logger.error("打包异常：%s", e)
            return PackageResult(success=False, message=f"打包过程异常：{str(e)[:200]}")

    # ------------------------------------------------------------------
    # zip — 高端定制 ZIP V5.0
    # ------------------------------------------------------------------
    def _build_zip(self) -> PackageResult:
        """V7.0 高端定制 ZIP 交付包 —— agent-spec 三层目录结构"""
        logger.info("开始 ZIP 打包：%s", self.script_path)

        safe_agent_name = self._sanitize_agent_name_for_zip()
        output_name = f"{safe_agent_name}_v1.0.zip"
        output_path = OUTPUT_DIR / output_name

        try:
            with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
                # 0. agent-meta.json —— 统一元数据（所有平台第一入口）
                try:
                    meta_json = self._build_agent_meta()
                    zf.writestr("agent-meta.json", meta_json)
                    logger.info("agent-meta.json 已注入")
                except Exception as meta_err:
                    logger.warning("agent-meta.json 生成失败: %s", meta_err)

                # 0b. prompt_pack —— 内核层 4 个 .md 文件
                try:
                    prompt_files = self._build_prompt_pack_files()
                    for arcname, content in prompt_files.items():
                        zf.writestr(arcname, content)
                        logger.info("%s 已注入", arcname)
                except Exception as pp_err:
                    logger.warning("prompt_pack 生成失败: %s", pp_err)

                # 1. 主程序 CLI 版
                zf.write(self.script_path, arcname="app.py")
                logger.info("app.py (CLI) 已注入")

                # 1b. Web 版
                if self.web_script_path and self.web_script_path.exists():
                    zf.write(self.web_script_path, arcname="app_web.py")
                    logger.info("app_web.py (Web) 已注入")

                # 2. 环境变量模板
                env_content = self._generate_env_file()
                zf.writestr("配置_API_Key.env", env_content)
                logger.info("配置_API_Key.env 已注入")

                # 3. requirements.txt
                requirements = ["openai>=1.0.0", "python-dotenv>=1.0.0", "streamlit>=1.28.0"]
                zf.writestr("requirements.txt", "\n".join(requirements) + "\n")

                # 4. 启动脚本
                bat_content, sh_content = self._generate_launch_scripts()
                zf.writestr("一键启动_Windows.bat", bat_content)
                zf.writestr("一键启动_Mac.sh", sh_content)
                logger.info("启动脚本已注入")

                # 4b. import.sh —— 一键导入脚本（自动检测平台）
                import_sh = self._generate_import_sh()
                zf.writestr("import.sh", import_sh)
                logger.info("import.sh 已注入")

                # 5. Word 说明书（主交付文档）
                docx_ok = False
                try:
                    docx_bytes = self._generate_docx_manual()
                    zf.writestr("使用说明书.docx", docx_bytes)
                    docx_ok = True
                    logger.info("使用说明书.docx 已注入")
                except Exception as docx_err:
                    logger.warning("docx 生成失败: %s", docx_err)

                # 6. Markdown 说明书（降级 or 纯文本补充）
                md_manual = self._generate_md_manual()
                if md_manual:
                    zf.writestr("使用说明书.md", md_manual)
                    label = "纯文本版本" if docx_ok else "降级方案"
                    logger.info("使用说明书.md 已注入 (%s)", label)

                # 7. 平台部署文件（adapters/ 子目录）
                platform_files = self._generate_platform_files()
                for arcname, content in platform_files.items():
                    zf.writestr(arcname, content)
                    logger.info("%s 已注入", arcname)

            logger.info("ZIP 打包完成 -> %s", output_path)
            return PackageResult(
                success=True,
                output_path=str(output_path),
                message=(
                    "交付包已生成\n"
                    "包含：agent-meta.json + prompts/ + adapters/ + import.sh + 主程序 + Word 说明书"
                ),
            )

        except Exception as e:
            logger.error("ZIP 打包异常：%s", e)
            return PackageResult(
                success=False,
                message=f"ZIP 打包失败：{str(e)[:200]}",
            )

    # ------------------------------------------------------------------
    # web — Streamlit 一键启动
    # ------------------------------------------------------------------
    def _build_web(self) -> PackageResult:
        zip_result = self._build_zip()
        if not zip_result.success:
            return zip_result
        return PackageResult(
            success=True,
            output_path=str(zip_result.output_path),
            message="Web 版交付包已生成（含 Streamlit 模板）",
        )

    # ------------------------------------------------------------------
    # docx — 仅导出说明书
    # ------------------------------------------------------------------
    def _build_docx_only(self) -> PackageResult:
        safe_name = self._sanitize_agent_name_for_zip()
        output_name = f"{safe_name}_说明书.docx"
        output_path = OUTPUT_DIR / output_name

        try:
            docx_bytes = self._generate_docx_manual()
            output_path.write_bytes(docx_bytes)
            logger.info("DOCX 说明书导出完成 -> %s", output_path)
            return PackageResult(
                success=True,
                output_path=str(output_path),
                message="Word 说明书已导出",
            )
        except Exception as e:
            logger.error("docx 导出失败：%s", e)
            return PackageResult(success=False, message=f"docx 导出失败：{str(e)[:200]}")

    # ==================================================================
    # 内部工具 — 文件生成器
    # ==================================================================

    def _sanitize_agent_name_for_zip(self) -> str:
        name = self.agent_name.strip()
        safe = re.sub(r"[^\w\u4e00-\u9fff\-_]", "_", name)
        safe = re.sub(r"_+", "_", safe).strip("_")
        return safe or "agent"

    def _generate_env_file(self) -> str:
        return (
            "# ============================================\n"
            "# API Key 配置文件\n"
            "# 使用方法：复制本文件，重命名为 .env，填入你的 Key 即可\n"
            "# ============================================\n\n"
            "MY_API_KEY=请填入你的API_Key\n"
            "MY_BASE_URL=https://api.deepseek.com/v1\n"
            "MY_MODEL_NAME=deepseek-chat\n"
        )

    def _generate_launch_scripts(self) -> tuple:
        """
        生成双击启动脚本（专业可用，自动处理工作目录）
        返回 (bat_content, sh_content)
        """
        safe_name = self.agent_name.replace("%", "%%").replace("^", "^^")

        bat_content = (
            "@echo off\n"
            "chcp 65001 >nul\n"
            f"title {safe_name}\n"
            "cd /d \"%~dp0\"\n"
            "echo ============================================\n"
            f"echo   {safe_name}\n"
            "echo ============================================\n"
            "echo.\n"
            "echo [1/3] 检查 Python 环境...\n"
            "python --version >nul 2>&1\n"
            "if %errorlevel% neq 0 (\n"
            "    echo [错误] 未检测到 Python，请先安装 Python 3.10+\n"
            "    echo 下载地址: https://www.python.org/downloads/\n"
            "    pause\n"
            "    exit /b 1\n"
            ")\n"
            "echo [通过] Python 已就绪\n"
            "echo.\n"
            "echo [2/3] 安装依赖...\n"
            "pip install -r requirements.txt -q -i https://pypi.tuna.tsinghua.edu.cn/simple\n"
            "if %errorlevel% neq 0 (\n"
            "    echo [警告] 清华源安装失败，尝试默认源...\n"
            "    pip install -r requirements.txt -q\n"
            ")\n"
            "echo [通过] 依赖安装完成\n"
            "echo.\n"
            "echo [提示] 请确保已将 配置_API_Key.env 复制为 .env 并填入 API Key\n"
            "echo.\n"
            "echo [3/3] 启动程序...\n"
            "python app.py\n"
            "echo.\n"
            "echo 运行结束，按任意键退出...\n"
            "pause\n"
        )

        sh_content = (
            "#!/usr/bin/env bash\n"
            "# 切换到脚本所在目录\n"
            'cd "$(dirname "$0")"\n'
            'echo "============================================"\n'
            f'echo "  {self.agent_name}"\n'
            'echo "============================================"\n'
            'echo ""\n'
            'echo "[1/3] 检查 Python 环境..."\n'
            'if command -v python3 &> /dev/null; then\n'
            '    PYTHON=python3\n'
            'elif command -v python &> /dev/null; then\n'
            '    PYTHON=python\n'
            'else\n'
            '    echo "[错误] 未检测到 Python，请先安装 Python 3.10+"\n'
            '    exit 1\n'
            'fi\n'
            'echo "[通过] $PYTHON 已就绪"\n'
            'echo ""\n'
            'echo "[2/3] 安装依赖..."\n'
            '$PYTHON -m pip install -r requirements.txt -q -i https://pypi.tuna.tsinghua.edu.cn/simple 2>/dev/null || $PYTHON -m pip install -r requirements.txt -q\n'
            'echo "[通过] 依赖安装完成"\n'
            'echo ""\n'
            'echo "[提示] 请确保已将 配置_API_Key.env 复制为 .env 并填入 API Key"\n'
            'echo ""\n'
            'echo "[3/3] 启动程序..."\n'
            '$PYTHON app.py\n'
            'echo ""\n'
            'echo "运行结束"\n'
        )

        return bat_content, sh_content

    # ------------------------------------------------------------------
    # 委托方法 — 委托给 builder 引擎（V7.0）
    # ------------------------------------------------------------------
    def _build_agent_meta(self) -> str:
        """委托 builder 生成 agent-meta.json"""
        try:
            from core.builder import AgentBuilder
            config = self._load_agent_config()
            if config:
                builder = AgentBuilder(config, delivery_type=self.delivery_type)
                return builder.build_agent_meta()
        except Exception as e:
            logger.warning("agent-meta 委托生成失败: %s", e)
        # 最小兜底
        return json.dumps({
            "name": self.agent_name,
            "version": "1.0.0",
            "platforms": ["coze", "dify"],
            "generated_at": datetime.now().isoformat(),
        }, ensure_ascii=False, indent=2)

    def _build_prompt_pack_files(self) -> dict:
        """委托 builder 生成 prompts/ 下 4 个 .md 文件"""
        try:
            from core.builder import AgentBuilder
            config = self._load_agent_config()
            if config:
                builder = AgentBuilder(config, delivery_type=self.delivery_type)
                return builder.build_prompt_pack_files()
        except Exception as e:
            logger.warning("prompt_pack 委托生成失败: %s", e)
        return {}

    def _generate_import_sh(self) -> str:
        """生成 import.sh — 一键导入脚本，自动检测平台并导入"""
        agent_name = self.agent_name
        return (
            "#!/usr/bin/env bash\n"
            "# ============================================================\n"
            f"#  {agent_name} — 一键导入脚本\n"
            "#  自动检测已安装的平台并导入适配配置\n"
            "# ============================================================\n"
            'set -e\n'
            'SCRIPT_DIR="$(cd "$(dirname "$0")" && pwd)"\n'
            'echo "=============================================="\n'
            f'echo "  {agent_name} — 平台导入工具"\n'
            'echo "=============================================="\n'
            'echo ""\n'
            '\n'
            '# --- 检测 Dify ---\n'
            'if [ -f "$SCRIPT_DIR/adapters/dify/agent.yaml" ]; then\n'
            '    echo "[dify] 适配配置已就绪: adapters/dify/agent.yaml"\n'
            '    echo "  导入方法: Dify 控制台 → 创建应用 → 导入 DSL → 选择此文件"\n'
            'else\n'
            '    echo "[dify] (未生成)"\n'
            'fi\n'
            '\n'
            '# --- 检测 Coze ---\n'
            'if [ -f "$SCRIPT_DIR/adapters/coze/bot.json" ]; then\n'
            '    echo "[coze] 适配配置已就绪: adapters/coze/bot.json"\n'
            '    echo "  导入方法: Coze 控制台 → 创建 Bot → 导入配置 → 选择此文件"\n'
            'else\n'
            '    echo "[coze] (未生成)"\n'
            'fi\n'
            '\n'
            '# --- 检测飞书 ---\n'
            'if [ -f "$SCRIPT_DIR/adapters/feishu/bot-config.yaml" ]; then\n'
            '    echo "[feishu] 适配配置已就绪: adapters/feishu/bot-config.yaml"\n'
            '    echo "  导入方法: 飞书开放平台 → 创建应用 → 按配置文件填写"\n'
            'else\n'
            '    echo "[feishu] (未生成)"\n'
            'fi\n'
            '\n'
            '# --- 检测 OpenClaw ---\n'
            'if [ -f "$SCRIPT_DIR/adapters/openclaw/config.yaml" ]; then\n'
            '    echo "[openclaw] 适配配置已就绪: adapters/openclaw/config.yaml"\n'
            '    echo "  导入方法: 复制到 OpenClaw 项目 agents/ 目录"\n'
            'else\n'
            '    echo "[openclaw] (未生成)"\n'
            'fi\n'
            '\n'
            'echo ""\n'
            'echo "导入完成。请参考各平台文档完成后续配置。"\n'
            'echo "元数据文件: agent-meta.json (所有适配层的统一入口)"\n'
        )

    # ------------------------------------------------------------------
    # 平台部署文件生成（V5.0 — 走 builder 引擎）
    # ------------------------------------------------------------------
    def _generate_platform_files(self) -> dict:
        try:
            from core.builder import AgentBuilder

            config = self._load_agent_config()
            if config:
                builder = AgentBuilder(config, delivery_type=self.delivery_type)
                return builder.build_platform_files()
            else:
                logger.warning("无法从脚本中读取 agent_config，使用简单替换")
                return self._generate_platform_files_simple()
        except Exception as e:
            logger.warning("平台文件生成器初始化失败，使用简单替换: %s", e)
            return self._generate_platform_files_simple()

    def _load_agent_config(self) -> dict:
        # 1. 同目录 agent_config.json
        config_path = self.script_path.parent / "agent_config.json"
        if config_path.exists():
            try:
                return json.loads(config_path.read_text(encoding="utf-8"))
            except Exception:
                pass

        # 2. 脚本内嵌 # AGENT_CONFIG: ...
        try:
            content = self.script_path.read_text(encoding="utf-8")
            match = re.search(r'# AGENT_CONFIG:\s*(\{.+\})', content, re.DOTALL)
            if match:
                return json.loads(match.group(1))
        except Exception:
            pass

        # 3. 回退最小配置
        return {
            "agent_name": self.agent_name,
            "system_prompt": self._get_system_prompt(),
            "delivery_type": self.delivery_type,
            "required_skills": [self.agent_scenario],
            "workflow_steps": [],
        }

    def _generate_platform_files_simple(self) -> dict:
        platform_files = {}
        platform_arcnames = {
            "feishu": "部署到飞书_机器人配置.yaml",
            "coze": "部署到Coze_导入文件.json",
            "dify": "部署到Dify_工作流导入.yml",
        }

        for platform_key, template_filename in PLATFORM_TEMPLATES.items():
            template_path = TEMPLATES_DIR / template_filename
            arcname = platform_arcnames[platform_key]

            if not template_path.exists():
                logger.warning("平台模板 %s 不存在，跳过", template_filename)
                continue

            content = template_path.read_text(encoding="utf-8")
            content = content.replace("{AGENT_NAME}", self.agent_name)
            content = content.replace("{AGENT_SCENARIO}", self.agent_scenario)
            escaped_prompt = self._escape_for_json(self._get_system_prompt())
            content = content.replace("{SYSTEM_PROMPT_ESCAPED}", escaped_prompt)

            platform_files[arcname] = content

        return platform_files

    def _get_system_prompt(self) -> str:
        try:
            content = self.script_path.read_text(encoding="utf-8")
            match = re.search(r'SYSTEM_PROMPT\s*=\s*"""(.+?)"""', content, re.DOTALL)
            if match:
                return match.group(1).strip()
        except Exception:
            pass
        return f"你是{self.agent_name}，专注于{self.agent_scenario}。"

    @staticmethod
    def _escape_for_json(text: str) -> str:
        return text.replace("\\", "\\\\").replace('"', '\\"').replace("\n", "\\n").replace("\r", "\\r").replace("\t", "\\t")

    # ------------------------------------------------------------------
    # Word 说明书（无 emoji，企业级交付）
    # ------------------------------------------------------------------
    def _generate_docx_manual(self) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("缺少 python-docx 依赖。请运行：pip install python-docx")

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)

        # 封面标题
        title = doc.add_heading(self.agent_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph("AI 智能代理 · 使用说明书")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        # 1. 产品概述
        doc.add_heading("1. 产品概述", level=1)
        doc.add_paragraph(f"本产品是一个名为「{self.agent_name}」的 AI 智能代理，专注于：")
        doc.add_paragraph(self.agent_scenario, style="List Bullet")
        doc.add_paragraph(
            "基于大语言模型（LLM），支持 DeepSeek、通义千问、智谱 GLM 以及任意兼容 OpenAI 接口的模型。"
            "用户仅需配置 API Key 即可开始使用。"
        )

        # 2. 文件清单
        doc.add_heading("2. 交付文件清单", level=1)
        doc.add_paragraph("解压后你将看到以下文件：")
        file_list = [
            ("app.py", "主程序，所有 AI 能力均在此"),
            ("配置_API_Key.env", "API Key 配置模板，复制为 .env 后填入 Key"),
            ("一键启动_Windows.bat", "Windows 双击启动"),
            ("一键启动_Mac.sh", "Mac / Linux 双击启动"),
            ("requirements.txt", "Python 依赖清单，启动脚本自动安装"),
            ("使用说明书.docx", "本文件，Word 格式正式说明书"),
            ("使用说明书.md", "Markdown 纯文本说明书（备选）"),
            ("部署到飞书_机器人配置.yaml", "飞书应用部署配置"),
            ("部署到Coze_导入文件.json", "Coze 平台导入配置"),
            ("部署到Dify_工作流导入.yml", "Dify 工作流 DSL"),
        ]
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "文件名"
        hdr_cells[1].text = "说明"
        for fname, desc in file_list:
            row_cells = table.add_row().cells
            row_cells[0].text = fname
            row_cells[1].text = desc
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

        # 3. 快速上手
        doc.add_heading("3. 快速上手（3 步）", level=1)

        doc.add_heading("3.1 安装 Python 依赖", level=2)
        doc.add_paragraph("方式一（推荐）：双击对应系统的「一键启动」脚本，自动完成依赖安装。")
        doc.add_paragraph("方式二：打开终端，进入本文件夹，执行以下命令：")
        p = doc.add_paragraph()
        run = p.add_run("pip install -r requirements.txt")
        run.font.name = "Courier New"
        run.font.size = Pt(10)

        doc.add_heading("3.2 配置 API Key", level=2)
        doc.add_paragraph(
            "找到「配置_API_Key.env」文件，将其复制一份并重命名为 .env 。"
            "用文本编辑器（如记事本）打开 .env ，按照里面的注释填入你的大模型 API Key。"
        )
        doc.add_paragraph("支持的模型与对应 Base URL：")
        model_list = [
            "DeepSeek: https://api.deepseek.com/v1",
            "通义千问: https://dashscope.aliyuncs.com/compatible-mode/v1",
            "智谱 GLM: https://open.bigmodel.cn/api/paas/v4",
            "其他兼容 OpenAI 接口: 你的服务商地址",
        ]
        for m in model_list:
            doc.add_paragraph(m, style="List Bullet")

        doc.add_heading("3.3 启动程序", level=2)
        doc.add_paragraph("Windows：双击「一键启动_Windows.bat」")
        doc.add_paragraph("Mac / Linux：终端运行 bash 一键启动_Mac.sh，或直接双击")
        doc.add_paragraph("手动启动：终端执行 python app.py")

        # 4. 多平台部署
        doc.add_heading("4. 多平台部署", level=1)
        doc.add_paragraph("本交付包包含三个平台的部署配置文件，可根据需要选用：")
        platforms = [
            "飞书：使用「部署到飞书_机器人配置.yaml」在飞书开放平台创建应用。简要步骤：创建应用 -> 配置机器人能力 -> 部署脚本到服务器 -> 配置回调地址 -> 发布上线。",
            "Coze：在 coze.com 创建 Bot 时选择导入「部署到Coze_导入文件.json」，填入 API Key 后发布，可一键绑定飞书、微信等渠道。",
            "Dify：在你的 Dify 实例中导入「部署到Dify_工作流导入.yml」DSL 文件，配置模型供应商后发布。",
        ]
        for plat in platforms:
            doc.add_paragraph(plat, style="List Bullet")

        # 5. 费用说明
        doc.add_heading("5. 费用说明", level=1)
        doc.add_paragraph(
            "本产品本身免费，仅需承担 AI 模型调用的 Token 费用。"
            "以 DeepSeek 官方价格为例，日常使用每月约数元至数十元。"
            "建议在 API 后台设置月消费上限，避免超额支出。"
        )

        # 6. 常见问题
        doc.add_heading("6. 常见问题", level=1)
        faqs = [
            ("Q: 双击启动后窗口闪退？",
             "A: 通常是 Python 未安装或未添加到系统 PATH。请打开终端执行 python --version 检查。如果未安装，请从 https://www.python.org/downloads/ 下载安装，安装时勾选「Add Python to PATH」。"),
            ("Q: 提示 ModuleNotFoundError: No module named 'openai'？",
             "A: 依赖未成功安装。在终端中进入本文件夹，执行 pip install -r requirements.txt。"),
            ("Q: 提示认证错误 (401/403)？",
             "A: 请检查 .env 文件的配置：1) API Key 是否正确填写；2) Base URL 是否与所用模型匹配。"),
            ("Q: 能否在手机或网页上使用？",
             "A: 可以。推荐将本产品部署到飞书或 Coze 平台，即可通过群聊或 Web 页面访问。"),
            ("Q: 支持本地部署的开源模型吗？",
             "A: 支持。将 .env 中的 MY_BASE_URL 改为你的本地模型服务地址即可。"),
        ]
        for q, a in faqs:
            p_q = doc.add_paragraph()
            run_q = p_q.add_run(q)
            run_q.bold = True
            doc.add_paragraph(a)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        doc.save(tmp_path)
        docx_bytes = Path(tmp_path).read_bytes()
        os.unlink(tmp_path)
        return docx_bytes

    def _generate_md_manual(self) -> str:
        """Markdown 说明书（纯文本版，无 emoji）"""
        return (
            f"# {self.agent_name} — 使用说明书\n\n"
            f"## 1. 产品概述\n\n"
            f"「{self.agent_name}」是一个 AI 智能代理，专注于：\n\n"
            f"{self.agent_scenario}\n\n"
            f"基于大语言模型，支持 DeepSeek / 通义千问 / 智谱 GLM / 任意 OpenAI 兼容接口。\n\n"
            f"## 2. 交付文件清单\n\n"
            f"| 文件名 | 说明 |\n"
            f"|--------|------|\n"
            f"| app.py | 主程序 |\n"
            f"| 配置_API_Key.env | API Key 配置模板（复制为 .env 后填入 Key）|\n"
            f"| 一键启动_Windows.bat | Windows 双击启动 |\n"
            f"| 一键启动_Mac.sh | Mac / Linux 启动脚本 |\n"
            f"| requirements.txt | Python 依赖清单 |\n"
            f"| 使用说明书.docx | Word 格式正式说明书 |\n"
            f"| 使用说明书.md | 本文件，纯文本说明书 |\n"
            f"| 部署到飞书_机器人配置.yaml | 飞书应用部署配置 |\n"
            f"| 部署到Coze_导入文件.json | Coze 平台导入配置 |\n"
            f"| 部署到Dify_工作流导入.yml | Dify 工作流 DSL |\n\n"
            f"## 3. 快速上手\n\n"
            f"### 3.1 安装依赖\n\n"
            f"方式一（推荐）：双击对应系统的「一键启动」脚本，自动安装。\n\n"
            f"方式二：在终端进入本文件夹，执行：\n\n"
            f"```bash\npip install -r requirements.txt\n```\n\n"
            f"### 3.2 配置 API Key\n\n"
            f"1. 将「配置_API_Key.env」复制一份，重命名为 .env\n"
            f"2. 用文本编辑器打开 .env，填入你的 API Key\n"
            f"3. 支持的模型与 Base URL：\n"
            f"   - DeepSeek: https://api.deepseek.com/v1\n"
            f"   - 通义千问: https://dashscope.aliyuncs.com/compatible-mode/v1\n"
            f"   - 智谱 GLM: https://open.bigmodel.cn/api/paas/v4\n"
            f"   - 其他: 你的服务商地址\n\n"
            f"### 3.3 启动\n\n"
            f"- Windows: 双击「一键启动_Windows.bat」\n"
            f"- Mac / Linux: 终端运行 `bash 一键启动_Mac.sh` 或直接双击\n"
            f"- 手动: `python app.py`\n\n"
            f"## 4. 多平台部署\n\n"
            f"- 飞书：使用「部署到飞书_机器人配置.yaml」在飞书开放平台创建应用\n"
            f"- Coze：在 coze.com 导入「部署到Coze_导入文件.json」\n"
            f"- Dify：导入 DSL 文件「部署到Dify_工作流导入.yml」\n\n"
            f"## 5. 费用说明\n\n"
            f"本产品免费，仅需承担 AI 模型调用的 Token 费用。\n"
            f"以 DeepSeek 为例，日常使用每月数元即可。建议设置消费上限。\n\n"
            f"## 6. 常见问题\n\n"
            f"**Q: 双击启动闪退？**\n"
            f"A: 检查 Python 是否安装且已添加到 PATH。终端执行 `python --version` 确认。\n\n"
            f"**Q: 提示 No module named 'openai'？**\n"
            f"A: 执行 `pip install -r requirements.txt` 安装依赖。\n\n"
            f"**Q: 提示认证错误？**\n"
            f"A: 检查 .env 文件中的 API Key 和 Base URL 是否正确。\n\n"
            f"**Q: 如何在手机上使用？**\n"
            f"A: 部署到飞书或 Coze 平台，即可通过手机访问。\n\n"
            f"---\n\n"
            f"*{self.agent_name} — 你的专属 AI 代理*\n"
        )
